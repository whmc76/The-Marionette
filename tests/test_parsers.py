"""Parser unit tests."""
import os
import pytest

ASSETS = os.path.join(os.path.dirname(__file__), "..", "assets")
LANTAO_BRIEF = os.path.join(ASSETS, "【0920】岚图梦想家评论brief.docx")
QIJING_BRIEF = os.path.join(ASSETS, "11月启境汽车品牌发布会&20日官方直播-尾翼.docx")


def _has_file(path: str) -> bool:
    return os.path.isfile(path)


@pytest.mark.skipif(not _has_file(LANTAO_BRIEF), reason="岚图brief not found")
def test_classified_parser_lantao():
    from src.parsers.classified import ClassifiedParser
    from src.parsers.base import load_paragraphs

    parser = ClassifiedParser()
    report = parser.parse(LANTAO_BRIEF)

    assert report.parser_name == "classified"
    assert report.confidence > 0
    assert report.spec is not None
    spec = report.spec
    assert spec.product_name, "product_name should be non-empty"
    assert len(spec.categories) >= 1, "should parse at least 1 category"


@pytest.mark.skipif(not _has_file(QIJING_BRIEF), reason="启境brief not found")
def test_script_lib_parser_qijing():
    from src.parsers.script_lib import ScriptLibParser

    parser = ScriptLibParser()
    report = parser.parse(QIJING_BRIEF)

    assert report.parser_name == "script_lib"
    assert report.spec is not None


def test_fallback_parser_with_docx():
    """FallbackParser should always return a valid spec."""
    import tempfile
    from docx import Document
    from src.parsers.fallback import FallbackParser

    # Create a minimal docx
    doc = Document()
    doc.add_paragraph("测试产品")
    doc.add_paragraph("这是一个非常好的产品，性能出色，使用体验极佳。")
    doc.add_paragraph("品质卓越，用户反馈非常正面。")

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name

    parser = FallbackParser()
    report = parser.parse(tmp_path)
    os.unlink(tmp_path)

    assert report.spec is not None
    assert report.confidence == pytest.approx(0.2)


@pytest.mark.skipif(not _has_file(LANTAO_BRIEF), reason="岚图brief not found")
def test_parse_brief_selects_best():
    from src.parsers import parse_brief

    best, all_reports = parse_brief(LANTAO_BRIEF)
    assert best.spec is not None
    assert len(all_reports) == 3  # classified + script_lib + fallback
    # Best should have highest confidence
    assert best.confidence == max(r.confidence for r in all_reports)


def test_llm_parser_json_extraction():
    from src.parsers.llm_parser import _extract_json, _extract_json_array

    # Plain object
    raw_obj = '{"product_name":"产品X","positive_ratio":0.6,"general_rules":["规则1"]}'
    data = _extract_json(raw_obj)
    assert data["product_name"] == "产品X"
    assert data["positive_ratio"] == pytest.approx(0.6)

    # Object wrapped in markdown fences
    raw_fenced = '```json\n' + raw_obj + '\n```'
    assert _extract_json(raw_fenced)["product_name"] == "产品X"

    # Plain array
    raw_arr = '[{"direction":"正向","theme":"使用体验","sub_themes":[],"description":"","personas":["真实用户"],"example_comments":["很好"]}]'
    arr = _extract_json_array(raw_arr)
    assert len(arr) == 1
    assert arr[0]["direction"] == "正向"

    # Array wrapped in an object — should pick list-of-dicts over list-of-strings
    raw_wrapped = '{"general_rules":["规则1","规则2"],"categories":[{"direction":"反击","theme":"竞品"}]}'
    arr2 = _extract_json_array(raw_wrapped)
    assert arr2[0]["direction"] == "反击"


def test_llm_parser_with_mock_client():
    import json, tempfile
    from unittest.mock import MagicMock
    from docx import Document
    from src.llm.base import LLMResponse
    from src.llm.client import LLMClient
    from src.parsers.llm_parser import LLMParser

    payload = {
        "title": "测试Brief",
        "product_name": "测试产品",
        "product_background": "一款优秀的测试产品",
        "general_rules": ["自然口语化"],
        "forbidden_phrases": ["最优惠"],
        "positive_ratio": 0.5,
        "min_char_length": 20,
        "platform_targets": ["微博"],
        "categories": [
            {"direction": "正向", "theme": "使用体验", "sub_themes": [], "description": "体验好",
             "personas": ["真实用户"], "example_comments": ["体验超好"]}
        ],
    }
    mock_backend = MagicMock()
    mock_backend.provider = "mock"
    mock_backend.chat.return_value = LLMResponse(
        content=json.dumps(payload, ensure_ascii=False),
        model="mock", provider="mock",
    )
    client = LLMClient(mock_backend, max_concurrency=1)

    doc = Document()
    doc.add_paragraph("测试产品评论brief")
    doc.add_paragraph("正向评论：使用体验非常出色")
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name

    parser = LLMParser(client)
    report = parser.parse(tmp_path)
    import os; os.unlink(tmp_path)

    assert report.parser_name == "llm"
    assert report.confidence > 0.7
    assert report.spec is not None
    assert report.spec.product_name == "测试产品"
    assert len(report.spec.categories) == 1


def test_parse_brief_prefers_llm_when_client_given():
    import json, tempfile
    from unittest.mock import MagicMock
    from docx import Document
    from src.llm.base import LLMResponse
    from src.llm.client import LLMClient
    from src.parsers import parse_brief

    payload = {
        "title": "T", "product_name": "LLM车", "product_background": "bg",
        "general_rules": [], "forbidden_phrases": [], "positive_ratio": 0.5,
        "min_char_length": 20, "platform_targets": [],
        "categories": [{"direction": "正向", "theme": "品质", "sub_themes": [],
                         "description": "", "personas": [], "example_comments": []}],
    }
    mock_backend = MagicMock()
    mock_backend.provider = "mock"
    mock_backend.chat.return_value = LLMResponse(
        content=json.dumps(payload), model="mock", provider="mock"
    )
    client = LLMClient(mock_backend, max_concurrency=1)

    doc = Document()
    doc.add_paragraph("产品brief")
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name

    best, all_reports = parse_brief(tmp_path, llm_client=client)
    import os; os.unlink(tmp_path)

    assert best.parser_name == "llm"
    assert best.spec.product_name == "LLM车"
    assert any(r.parser_name == "llm" for r in all_reports)


def test_load_paragraphs_with_tables():
    """load_paragraphs should extract table cell text."""
    import tempfile
    from docx import Document
    from src.parsers.base import load_paragraphs

    doc = Document()
    doc.add_paragraph("标题段落")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "单元格A"
    table.cell(0, 1).text = "单元格B"
    table.cell(1, 0).text = "单元格C"
    table.cell(1, 1).text = "单元格D"

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name

    paras = load_paragraphs(tmp_path)
    os.unlink(tmp_path)

    assert "标题段落" in paras
    assert "单元格A" in paras
