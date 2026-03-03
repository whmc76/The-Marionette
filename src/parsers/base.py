"""Parser base class and shared utilities."""
from __future__ import annotations

import re
from abc import ABC, abstractmethod
from typing import Optional

from docx import Document

from src.models import BriefSpec, ParseReport


def load_paragraphs(path: str) -> list[str]:
    """Load all non-empty paragraphs from a DOCX file."""
    doc = Document(path)
    paras = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            paras.append(text)
    # Also pull text from table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = p.text.strip()
                    if text and text not in paras:
                        paras.append(text)
    return paras


def clean_text(s: str) -> str:
    """Strip leading punctuation/whitespace markers."""
    return re.sub(r"^[\s\-\u2022\u25cf\uff65\u30fb\u00b7]+", "", s).strip()


class BaseParser(ABC):
    """Abstract base for all brief parsers."""

    name: str = "base"

    def parse(self, path: str) -> ParseReport:
        paragraphs = load_paragraphs(path)
        spec, confidence, warnings = self._parse(paragraphs)
        return ParseReport(
            parser_name=self.name,
            confidence=confidence,
            category_count=len(spec.categories) if spec else 0,
            persona_count=len({p for cat in (spec.categories if spec else []) for p in cat.personas}),
            forbidden_phrase_count=len(spec.forbidden_phrases) if spec else 0,
            example_count=sum(len(c.example_comments) for c in (spec.categories if spec else [])),
            warnings=warnings,
            spec=spec,
        )

    @abstractmethod
    def _parse(self, paragraphs: list[str]) -> tuple[Optional[BriefSpec], float, list[str]]:
        """Return (spec, confidence 0-1, warnings)."""
        ...
